import json
import os
from tempfile import NamedTemporaryFile as nft
from subprocess import run as runcommand

from PIL import Image as img
from docx.image.exceptions import UnrecognizedImageError
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Cm, Pt
from docx import Document
import yaml

from netwk import get_full_hmz, decode_add
from config import CONFIG, save_json, read_json, LANG

ldr = os.listdir
ext = os.path.exists
FONT_SIZE = [18, 16, 15, 14]

ENABLE_PANDOC: bool = CONFIG['ENABLE_PANDOC']
COMPLICATE_SLICE_NAME: bool = CONFIG["COMPLICATE_SLICE_NAME"]


def makedir(fname):
    if not os.path.exists(fname):
        os.mkdir(fname)


def find_hmz(path):
    """

    :param path:
    :return: Only the filename without global path.
    """
    for i in ldr(path):
        if '.hmz' in i:
            return i
    return 0


def read_hmz(filepath, complicated=False, source=False):
    try:
        hmz = read_json(filepath)
    except (json.JSONDecodeError, UnicodeDecodeError):
        numnm = filepath.split('/')[-1][:-4]
        hmz = get_full_hmz(numnm, filepath)
        save_json(filepath, hmz)
    if source:
        return hmz
    if not complicated:
        return hmz['name'], hmz['writer'], hmz['allnet'], hmz['allname']
    return hmz['name'], hmz['writer'], hmz['allnet'], hmz['allname'], hmz['bunko'], hmz['discription']


# noinspection PyProtectedMember
def put_headings(document, content, lev):
    content = content.strip()
    heading = document.add_heading(content, level=lev + 1)
    fs = FONT_SIZE[lev]
    run = heading.runs[0]
    run.font.name_t = '宋体'
    run.font.size = Pt(fs)
    run.font.color.rgb = RGBColor(0, 0, 0)

    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)

    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(fs / 2)
    heading_format.space_after = Pt(fs / 2)


# noinspection PyProtectedMember
def set_font_to(run, fontname):
    r = run._element
    rPr = r.find(qn('w:rPr'))
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), fontname)
    rPr.append(rFonts)


def put_title_page(docx, title, author, description=None):
    docx.add_paragraph('\n\n\n\n\n\n')

    title_paragraph = docx.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(20)
    set_font_to(title_run, '宋体')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    docx.add_paragraph('\n\n\n\n')

    author_paragraph = docx.add_paragraph()
    author_run = author_paragraph.add_run(author)
    author_run.font.size = Pt(16)
    author_run.font.name = '宋体'
    set_font_to(author_run, '宋体')
    author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if description:
        docx.add_paragraph('\n\n\n\n\n')
        description_paragraph = docx.add_paragraph()
        description_run = description_paragraph.add_run("    简介：")
        description_run.font.size = Pt(12)
        description_run.font.name = '宋体'
        set_font_to(description_run, '宋体')

        description_paragraph.add_run(description)
        description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    docx.add_page_break()


def get_size(picture):
    ig = img.open(picture)
    m = ig.size
    ig.close()
    return m


def put_pict(doc, picture):
    siz = get_size(picture)
    if siz[1] / siz[0] > 1.5:
        doc.add_picture(picture, height=Cm(22.8))
        return
    doc.add_picture(picture, width=Cm(15.2))


def ordered_ldr(path, typ='.txt'):
    pre_ldrlist = ldr(path)
    b_ill = False
    if '插图' in pre_ldrlist:
        pre_ldrlist.remove('插图')
        b_ill = True
    ldrlist = []
    for i in pre_ldrlist:
        try:
            ldrlist.append(int(i.split('.')[0]))
        except ValueError:
            continue
    ldrlist.sort()
    ldrlist = [str(i) + typ for i in ldrlist]
    return ldrlist, b_ill


def ordered_volumes(path, all_):
    g_vols = [i[0] for i in all_]
    hf = ldr(path)
    vols = []
    for i in g_vols:
        if i in hf:
            vols.append(i)
            continue
        return vols
    return vols


def Novel_docx():
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "SimSun"
    normal_style.font.size = Pt(12)
    rPr = normal_style.element.get_or_add_rPr()
    rFonts = rPr.rFonts if rPr.rFonts is not None else OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rPr.append(rFonts)
    pPr = normal_style.element.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), '480')
    pPr.append(ind)

    return doc


def get_cover_from(volume):
    if '插图' not in ldr(volume):
        return ''
    return volume + '/插图/' + ordered_ldr(f'{volume}/插图', '.jpg')[0][0]


class NotAHFolderError(Exception):
    def __init__(self):
        print('Not A HFolder')


class HFolder(object):
    def __init__(self, folder_adr):
        self.folder = folder_adr

        adr_slice = folder_adr.split('/')
        self.name = adr_slice[-1]
        self.globa = '/'.join(adr_slice[:-1]) + '/'

        self.hmz = find_hmz(folder_adr)
        self.numname = self.hmz.split('.')[0] if self.hmz else ''
        if ENABLE_PANDOC:
            hmzfile = read_hmz(self.folder + '/' + self.hmz, complicated=True)
            self.name, self.writer, self.allnet, self.allname, self.bunko, self.discription = hmzfile
        else:
            hmzfile = read_hmz(self.folder + '/' + self.hmz)
            self.name, self.writer, self.allnet, self.allname = hmzfile
            self.bunko = ''
            self.discription = ''

        self.volumes = [i[0] for i in self.allnet]
        dfc = get_cover_from(f'{self.globa}{self.name}/{self.volumes[0]}')
        self.cover = dfc if ext(dfc) else None

    @staticmethod
    def format_pic_fdr(path):
        for i in ldr(path):
            imgg = img.open(path + f'/{i}')
            imgg.save(path + f'/{i}', "JPEG", quality=100, optimize=True)
            imgg.close()

    @staticmethod
    def form_text(lines, docx):
        m = lines.pop(0)
        put_headings(docx, m, 1)
        for i in lines:
            docx.add_paragraph(i.strip())

    @staticmethod
    def form_pict_fdr(path, docx):
        path_file = ordered_ldr(path, '.jpg')[0]
        for i in path_file:
            try:
                put_pict(docx, path + '/' + i)
            except UnrecognizedImageError:
                HFolder.format_pic_fdr(path)
                put_pict(docx, path + '/' + i)
        return path_file

    def form_each_volume(self, chapdir):
        cover = None
        doc_each = Document()
        put_title_page(doc_each, self.name, self.writer)
        vname = chapdir.split('/')[-1]
        put_headings(doc_each, vname, 0)

        for i in range(len(ldr(chapdir))):
            try:
                f = open(f'{chapdir}/{i}.txt', 'r', encoding='utf-8')
                fl = f.readlines()
                self.form_text(fl, doc_each)
            except FileNotFoundError:
                continue

        if '插图' in ldr(chapdir):
            put_headings(doc_each, ' 插图', 1)
            self.form_pict_fdr(chapdir + '/插图', doc_each)
            cover = chapdir + '/插图/1.jpg'

        return cover, doc_each

    def formfd(self, mode=1, complicate_name: bool = COMPLICATE_SLICE_NAME):
        if mode == 1:
            self.formslice([], self.name, False)
            return 0

        cov_mat = []
        makedir(self.folder + '/Volume_docx')
        for i in self.allnet:
            try:
                cov, d = self.form_each_volume(self.folder + '/' + i[0])
                dname = f'{self.folder}/Volume_docx/{i[0]}.docx'
                if complicate_name:
                    dname = f'{self.folder}/Volume_docx/{self.name}_{i[0]}.docx'
                cov_mat.append((dname, cov))
                d.save(dname)
            except NotADirectoryError or FileNotFoundError:
                continue
        self.cover = cov_mat
        return 0

    def formepub(self, mode=1):
        if mode == 1:
            inpu = f'{self.folder}/{self.name}.docx'
            oupu = f'{self.folder}/{self.name}.epub'
            if ENABLE_PANDOC:
                convert_to_epub_pandoc(inpu, oupu, self.name, self.writer, self.numname, self.bunko,
                                       cover=self.cover, discription=self.discription)
                return 0
            convert_to_epub(inpu, oupu, self.name, writer=self.writer, cover=self.cover)
            return 0

        makedir(self.folder + '/Volume_epub')
        for i, cov in self.cover:
            fname = i.split('/')[-1][:-5]
            oupu = f'{self.folder}/Volume_epub/{fname}.epub'
            if ENABLE_PANDOC:
                convert_to_epub_pandoc(i, oupu, fname, writer=self.writer, cover=cov, numname=self.numname)
                continue
            convert_to_epub(i, oupu, fname, writer=self.writer, cover=cov)
        return None

    def formazw3(self, mode=1):
        if mode == 1:
            convert_to_azw3(self.folder + '/' + self.name + '.epub', self.folder + '/' + self.name + '.azw3')
            return 0

        makedir(self.folder + '/Volume_azw3')
        for i, cov in self.cover:
            slicename = i.split('/')
            smallname = slicename[-1][:-5]
            fname = '/'.join(slicename[:-2]) + '/Volume_epub/' + smallname + '.epub'
            convert_to_azw3(fname, self.folder + '/Volume_azw3/' + fname + '.azw3', fname)
        return None

    def formslice(self, slice_ind, filename, slice_=True):
        goal_doc = Novel_docx()
        goalname = self.folder + '/' + filename
        if slice_:
            makedir(self.folder + '/Slices')
            goalname = self.folder + '/Slices/' + filename

        partial_allnet = []
        for i in range(len(self.allnet)):
            if not slice_ind:
                partial_allnet = self.allnet
                break
            if slice_ind[i]:
                partial_allnet.append(self.allnet[i])

        put_title_page(goal_doc, self.name, self.writer)
        l_p_a = len(partial_allnet)
        for i in range(l_p_a):
            progress = (i + 1) / l_p_a * 100
            filled_length = int(30 * (i + 1) // l_p_a)
            bar = '=' * filled_length + ' ' * (30 - filled_length)
            print(f"FM: [{bar}] {progress:.2f}% ({i + 1}/{l_p_a})", end='\r')

            vol_name = partial_allnet[i][0]
            i_vol = len(vol_name.split(' '))
            put_headings(goal_doc, vol_name, 0)
            curdir = self.folder + '/' + vol_name
            curcontent, b_ill = ordered_ldr(curdir)
            for j in curcontent:
                chappath = curdir + '/' + j
                chap = open(chappath, 'r', encoding='utf-8')
                lines = chap.readlines()
                chap.close()
                lines[0] = ' '.join(lines[0].split(' ')[i_vol:])
                self.form_text(lines, goal_doc)
            if not b_ill:
                continue
            put_headings(goal_doc, f'插图', 1)
            self.form_pict_fdr(curdir + '/插图', goal_doc)
        print(LANG['TX_Finished'])
        goal_doc.save(goalname + '.docx')
        return 1


def convert_to_epub(pt, final_pt, title, writer, cover=''):
    command = ['ebook-convert', pt, final_pt, '--title', title, '--level1-toc', '//h:h1', '--level2-toc', '//h:h2']
    command = command + ['--authors', writer] if writer else command
    command = command + ['--cover', cover] if cover else command
    print(command)
    runcommand(command)


def convert_to_azw3(pt, final_pt, author='', cover=''):
    command = ['ebc/ebook-convert', pt, final_pt]
    command = command + ['--authors', author] if author else command
    command = command + ['--cover', cover] if cover else command
    print(command)
    runcommand(command)


def generate_yaml(name, writer, bunko, cover=None, discription=None):
    metadata = {'title': [{'type': 'main', 'text': name}],
                'creator': [{'role': 'writer', 'text': writer}],
                'publisher': bunko,
                'cover-image': cover,
                'description': discription}
    with nft("w", encoding="utf-8", delete=False) as y:
        y.write('---\n')
        yaml.dump(metadata, y, allow_unicode=True, indent=2, sort_keys=False, default_flow_style=False)
        y.write('...')
        return y.name


def convert_to_epub_pandoc(input_, output_, name, writer, numname, *args, cover: str = None, **kwargs):
    bunko = args[0] if args else None
    discription = kwargs.get('discription', None)
    if cover is None:
        cover = f'images/thumbnails/{numname}.jpg'
    if cover == '':
        cover = 'images/thumbnails/uncovered.jpg'
    print(name, writer, bunko, cover, discription, sep='\n-')
    y = generate_yaml(name, writer, bunko, cover, discription)
    command = ['pandoc', input_, '-o', output_, '--epub-title-page=false',
               f'--metadata-file={y}', f'--css={'appending.css'}']
    runcommand(command)
    try:
        os.unlink(y)
    finally:
        pass


def ui_2_py(ui, py):
    command = ['pyuic6', '-x', f'ui/{ui}', '-o', py]
    runcommand(command)


def activate():
    doc = Document()
    put_title_page(doc, "我的标题", "作者姓名")
    doc.save('C:/Users/Administrator/Desktop/1122.docx')


def activate2():
    convert_to_epub_pandoc('D:/ACGN/Novel/Re：从零开始的异世界生活/Re：从零开始的异世界生活.docx',
                           'D:/ACGN/Novel/Re：从零开始的异世界生活/Re：从零开始的异世界生活.epub',
                           'Re: 从零开始的异世界生活', '长月达平', 'MF文库J',
                           cover='D:/Program Files/LightND/images/thumbnails/1861.jpg',
                           discription='走出便利商店要回家的高中生‧菜月昴突然被召唤到异世界。\n'
                                       '这莫非就是很流行的异世界召唤!?可是眼前没有召唤者就算了，还遭遇强盗迅速面临性命危机。'
                                       '\n这时，一名神秘银发美少女和猫精灵拯救了一筹莫展的他。'
                                       '\n以报恩为名义，昴自告奋勇要帮助少女找东西。'
                                       '\n然而，好不容易才掌握到线索，昴和少女却被不明人士攻击而殒命──本来应该是这样，但回过神来，昴却发现自己置身在第一次被召唤到这个异世界时的所在位置。'
                                       '\n「死亡回归」──无力的少年得到的唯一能力，是死后时间会倒转回到一开始。跨越无数绝望，从死亡的命运中拯救少女！')


def activate3():
    convert_to_azw3("D:/HuaweiMoveData/Users/he660/Desktop/test/穿越时空的约定.epub", 'D:/HuaweiMoveData/Users/he660/Desktop/test/穿越时空的约定.azw3', 'mr')


def activate4():
    print(decode_add('D:/ACGN/Novel/无职转生～到了异世界就拿出真本事～(无职转生~在异世界认真地活下去~)/第一卷 幼年期/插图/66110.jpg'))


if __name__ == '__main__':
    activate4()
